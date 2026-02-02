"""
Evacuatie Calculator - Streamlit Web Applicatie
BBL Art. 4.80/4.81 - Opvang- en doorstroomcapaciteit trappenhuizen

Volledig werkende web versie van de desktop applicatie.
"""

import streamlit as st
import plotly.graph_objects as go
import plotly.express as px
from datetime import date
import json
import io
import google.generativeai as genai
from fpdf import FPDF
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import pandas as pd

# Import bestaande modules
from models.constanten import (
    DoorgangType, LocatieType, VerlaatTijdType,
    ONTRUIMINGSTIJDEN, VERLAAT_TIJDEN,
    STANDAARD_HOOGTEVERSCHIL, STANDAARD_TRAPBREEDTE,
    STANDAARD_AANTAL_TREDEN, STANDAARD_BORDES_OPP,
    STANDAARD_TUSSENBORDES_OPP, STANDAARD_DOORGANG_BREEDTE,
)
from models.trap import Trap, TrapVerdieping
from models.project import Project, Verdieping
from berekeningen.simulatie import SimulatieEngine, simuleer_alle_trappen
from berekeningen.toetsing import toets_alle_resultaten, ToetsStatus

# Pagina configuratie
st.set_page_config(
    page_title="Evacuatie Calculator - BBL 4.80/4.81",
    page_icon="🏢",
    layout="wide",
    initial_sidebar_state="expanded"
)

# Custom CSS voor professionele uitstraling
st.markdown("""
<style>
    .main-header {
        background: linear-gradient(135deg, #1565C0, #0D47A1);
        color: white;
        padding: 1.5rem 2rem;
        border-radius: 10px;
        margin-bottom: 2rem;
    }
    .main-header h1 {
        margin: 0;
        font-size: 2rem;
    }
    .main-header p {
        margin: 0.5rem 0 0 0;
        opacity: 0.9;
    }
    .stMetric {
        background-color: #f8f9fa;
        padding: 1rem;
        border-radius: 8px;
        border-left: 4px solid #1565C0;
    }
    .success-box {
        background-color: #d4edda;
        border: 1px solid #c3e6cb;
        border-left: 4px solid #28a745;
        padding: 1rem;
        border-radius: 4px;
        margin: 1rem 0;
    }
    .warning-box {
        background-color: #fff3cd;
        border: 1px solid #ffeeba;
        border-left: 4px solid #ffc107;
        padding: 1rem;
        border-radius: 4px;
        margin: 1rem 0;
    }
    .error-box {
        background-color: #f8d7da;
        border: 1px solid #f5c6cb;
        border-left: 4px solid #dc3545;
        padding: 1rem;
        border-radius: 4px;
        margin: 1rem 0;
    }
    .info-box {
        background-color: #e7f3ff;
        border: 1px solid #b8daff;
        border-left: 4px solid #1565C0;
        padding: 1rem;
        border-radius: 4px;
        margin: 1rem 0;
    }
    div[data-testid="stSidebar"] {
        background-color: #f8f9fa;
    }
    .step-indicator {
        display: flex;
        justify-content: space-between;
        margin-bottom: 2rem;
    }
    .step {
        flex: 1;
        text-align: center;
        padding: 0.5rem;
        border-bottom: 3px solid #e0e0e0;
    }
    .step.active {
        border-bottom-color: #1565C0;
        font-weight: bold;
    }
    .step.completed {
        border-bottom-color: #28a745;
    }
</style>
""", unsafe_allow_html=True)


def init_session_state():
    """Initialiseer session state variabelen"""
    if 'stap' not in st.session_state:
        st.session_state.stap = 1

    if 'project' not in st.session_state:
        st.session_state.project = {
            'projectnaam': 'Nieuw Project',
            'datum': date.today(),
            'aantal_trappen': 1,
            'aantal_bouwlagen': 10,
            'laagste_verdieping': 0,
            'voorportalen': False,
        }

    if 'personen' not in st.session_state:
        st.session_state.personen = {}

    if 'trappen' not in st.session_state:
        st.session_state.trappen = {}

    if 'resultaten' not in st.session_state:
        st.session_state.resultaten = None

    if 'toetsingen' not in st.session_state:
        st.session_state.toetsingen = None


def render_header():
    """Render de header met titel en stappen indicator"""
    st.markdown("""
    <div class="main-header">
        <h1>🏢 Evacuatie Calculator</h1>
        <p>BBL Art. 4.80/4.81 - Opvang- en doorstroomcapaciteit trappenhuizen</p>
    </div>
    """, unsafe_allow_html=True)


def render_sidebar():
    """Render de sidebar met navigatie"""
    with st.sidebar:
        st.markdown("### 📋 Navigatie")

        stappen = ["Project", "Personen", "Trappen", "Resultaten"]

        for i, stap_naam in enumerate(stappen, 1):
            if i < st.session_state.stap:
                icon = "✅"
            elif i == st.session_state.stap:
                icon = "▶️"
            else:
                icon = "⬜"

            if st.button(f"{icon} {i}. {stap_naam}", key=f"nav_{i}",
                        disabled=(i > st.session_state.stap + 1),
                        width="stretch"):
                st.session_state.stap = i
                st.rerun()

        st.markdown("---")
        st.markdown("### ℹ️ BBL Referentie")
        st.markdown("""
        - **Art. 4.80**: Doorstroomcapaciteit
        - **Art. 4.81**: Opvangcapaciteit
        - **Lid 1**: Ontruimingstijd (15/20/30 min)
        - **Lid 2**: Compartiment < 1 min
        - **Lid 3**: Verdieping < 3,5 min
        """)

        st.markdown("---")
        st.markdown("### 📥 Project")

        # Download project als JSON
        if st.button("💾 Project opslaan", width="stretch"):
            project_data = {
                'project': st.session_state.project,
                'personen': st.session_state.personen,
                'trappen': st.session_state.trappen,
            }
            # Convert date to string
            project_data['project']['datum'] = str(project_data['project']['datum'])

            json_str = json.dumps(project_data, indent=2, ensure_ascii=False)
            st.download_button(
                label="⬇️ Download JSON",
                data=json_str,
                file_name=f"{st.session_state.project['projectnaam']}.json",
                mime="application/json",
                width="stretch"
            )

        # Upload project
        uploaded_file = st.file_uploader("📂 Project laden", type=['json'])
        if uploaded_file is not None:
            try:
                data = json.load(uploaded_file)
                st.session_state.project = data['project']
                st.session_state.project['datum'] = date.fromisoformat(data['project']['datum'])
                st.session_state.personen = data.get('personen', {})
                st.session_state.trappen = data.get('trappen', {})
                st.success("Project geladen!")
                st.rerun()
            except Exception as e:
                st.error(f"Fout bij laden: {e}")


def render_stap1_project():
    """Render stap 1: Projectgegevens"""
    st.markdown("## 📁 Stap 1: Projectgegevens")

    col1, col2 = st.columns(2)

    with col1:
        st.markdown("### Projectinformatie")

        st.session_state.project['projectnaam'] = st.text_input(
            "Projectnaam",
            value=st.session_state.project['projectnaam'],
            placeholder="Bijv. Kantoorgebouw Amsterdam"
        )

        st.session_state.project['datum'] = st.date_input(
            "Datum berekeningen",
            value=st.session_state.project['datum']
        )

    with col2:
        st.markdown("### Gebouwconfiguratie")

        st.session_state.project['aantal_trappen'] = st.number_input(
            "Aantal trappenhuizen",
            min_value=1,
            max_value=10,
            value=st.session_state.project['aantal_trappen'],
            help="Aantal trappenhuizen in het gebouw"
        )

        st.session_state.project['aantal_bouwlagen'] = st.number_input(
            "Aantal bouwlagen",
            min_value=1,
            max_value=60,
            value=st.session_state.project['aantal_bouwlagen'],
            help="Totaal aantal bouwlagen (inclusief begane grond)"
        )

        st.session_state.project['laagste_verdieping'] = st.number_input(
            "Laagste verdieping",
            min_value=-10,
            max_value=0,
            value=st.session_state.project['laagste_verdieping'],
            help="0 = begane grond, -1 = kelder"
        )

        st.session_state.project['voorportalen'] = st.checkbox(
            "Voorportalen aanwezig",
            value=st.session_state.project['voorportalen'],
            help="Zijn er voorportalen tussen verdieping en trappenhuis?"
        )

    # Info box
    hoogste = st.session_state.project['laagste_verdieping'] + st.session_state.project['aantal_bouwlagen'] - 1
    st.markdown(f"""
    <div class="info-box">
        <strong>Samenvatting:</strong><br>
        Verdiepingen: {st.session_state.project['laagste_verdieping']} t/m {hoogste}<br>
        Aantal trappenhuizen: {st.session_state.project['aantal_trappen']}
    </div>
    """, unsafe_allow_html=True)

    # Navigatie
    col1, col2, col3 = st.columns([1, 2, 1])
    with col3:
        if st.button("Volgende ➡️", type="primary", width="stretch"):
            # Initialiseer personen als nodig
            laagste = st.session_state.project['laagste_verdieping']
            hoogste = laagste + st.session_state.project['aantal_bouwlagen'] - 1
            for v in range(hoogste, laagste - 1, -1):
                if str(v) not in st.session_state.personen:
                    st.session_state.personen[str(v)] = 0

            st.session_state.stap = 2
            st.rerun()


def render_stap2_personen():
    """Render stap 2: Personen per verdieping"""
    st.markdown("## 👥 Stap 2: Personen per Verdieping")

    laagste = st.session_state.project['laagste_verdieping']
    hoogste = laagste + st.session_state.project['aantal_bouwlagen'] - 1

    # Initialiseer reset counter als die niet bestaat
    if 'personen_reset_counter' not in st.session_state:
        st.session_state.personen_reset_counter = 0

    st.markdown("""
    <div class="info-box">
        Voer het aantal personen in per verdieping.
        De begane grond (verdieping 0) telt meestal niet mee omdat mensen daar direct naar buiten kunnen.
    </div>
    """, unsafe_allow_html=True)

    # Snelle invoer opties
    st.markdown("### ⚡ Snelle invoer")
    col1, col2, col3 = st.columns(3)

    with col1:
        default_personen = st.number_input("Standaard per verdieping", min_value=0, value=50, key="default_pers")
    with col2:
        if st.button("Toepassen op alle verdiepingen"):
            for v in range(hoogste, laagste - 1, -1):
                if v != 0:  # Skip begane grond
                    st.session_state.personen[str(v)] = default_personen
            st.session_state.personen_reset_counter += 1
            st.rerun()
    with col3:
        if st.button("Reset naar 0"):
            for v in range(hoogste, laagste - 1, -1):
                st.session_state.personen[str(v)] = 0
            st.session_state.personen_reset_counter += 1
            st.rerun()

    st.markdown("---")
    st.markdown("### 📊 Personen per verdieping")

    # Maak kolommen voor invoer
    cols = st.columns(4)

    # Gebruik counter in key zodat widgets opnieuw worden aangemaakt na reset
    counter = st.session_state.personen_reset_counter

    for i, v in enumerate(range(hoogste, laagste - 1, -1)):
        col_idx = i % 4
        with cols[col_idx]:
            label = f"Verdieping {v}" if v != 0 else "Begane grond (0)"
            st.session_state.personen[str(v)] = st.number_input(
                label,
                min_value=0,
                max_value=1000,
                value=st.session_state.personen.get(str(v), 0),
                key=f"pers_{v}_{counter}"
            )

    # Totaal
    totaal = sum(st.session_state.personen.values())
    st.markdown(f"""
    <div class="info-box">
        <strong>Totaal aantal personen: {totaal}</strong>
    </div>
    """, unsafe_allow_html=True)

    # Navigatie
    col1, col2, col3 = st.columns([1, 2, 1])
    with col1:
        if st.button("⬅️ Terug", width="stretch"):
            st.session_state.stap = 1
            st.rerun()
    with col3:
        if st.button("Volgende ➡️", type="primary", width="stretch"):
            if totaal == 0:
                st.warning("Voer minimaal 1 persoon in!")
            else:
                # Initialiseer trappen
                init_trappen()
                st.session_state.stap = 3
                st.rerun()


def init_trappen():
    """Initialiseer trap configuratie"""
    trap_letters = "ABCDEFGHIJ"
    laagste = st.session_state.project['laagste_verdieping']
    hoogste = laagste + st.session_state.project['aantal_bouwlagen'] - 1

    for i in range(st.session_state.project['aantal_trappen']):
        letter = trap_letters[i] if i < len(trap_letters) else str(i + 1)
        trap_key = f"Trap {letter}"

        if trap_key not in st.session_state.trappen:
            st.session_state.trappen[trap_key] = {
                'locatie': 'BESCHERMD',
                'verlaat_tijd': 'STANDAARD',
                'doorgang_uitgang': 0.9,
                'type_uitgang': 'ENKELE_DEUR_LT135',
                'verdiepingen': {}
            }

            # Initialiseer verdiepingen
            for v in range(hoogste, laagste - 1, -1):
                st.session_state.trappen[trap_key]['verdiepingen'][str(v)] = {
                    'doorgang_trap': STANDAARD_DOORGANG_BREEDTE,
                    'type_doorgang': 'ENKELE_DEUR_LT135',
                    'bordes': STANDAARD_BORDES_OPP,
                    'tussenbordes': STANDAARD_TUSSENBORDES_OPP,
                    'hoogteverschil': STANDAARD_HOOGTEVERSCHIL,
                    'trapbreedte': STANDAARD_TRAPBREEDTE,
                    'treden': STANDAARD_AANTAL_TREDEN,
                }


def render_stap3_trappen():
    """Render stap 3: Trap configuratie"""
    st.markdown("## 🚪 Stap 3: Trap Configuratie")

    trap_namen = list(st.session_state.trappen.keys())

    if not trap_namen:
        init_trappen()
        trap_namen = list(st.session_state.trappen.keys())

    # Selecteer trap
    geselecteerde_trap = st.selectbox(
        "Selecteer trappenhuis",
        trap_namen,
        key="trap_select"
    )

    trap_data = st.session_state.trappen[geselecteerde_trap]

    # Trap algemeen
    st.markdown("### 🏢 Trappenhuis algemeen")
    col1, col2 = st.columns(2)

    with col1:
        locatie_opties = {
            'VEILIGHEIDSVLUCHTROUTE': 'Veiligheidsvluchtroute (30 min)',
            'EXTRA_BESCHERMD': 'Extra beschermde vluchtroute (20 min)',
            'BESCHERMD': 'Beschermde vluchtroute (15 min)',
        }
        trap_data['locatie'] = st.selectbox(
            "Type vluchtroute",
            list(locatie_opties.keys()),
            format_func=lambda x: locatie_opties[x],
            index=list(locatie_opties.keys()).index(trap_data.get('locatie', 'BESCHERMD')),
            key=f"locatie_{geselecteerde_trap}"
        )

        verlaat_opties = {
            'STANDAARD': 'Standaard (3,5 min)',
            'VERLENGD': 'Verlengd (6 min) - WBDBO>=30 + R200',
        }
        trap_data['verlaat_tijd'] = st.selectbox(
            "Max tijd verlaten verdieping (BBL 4.81 lid 3)",
            list(verlaat_opties.keys()),
            format_func=lambda x: verlaat_opties[x],
            index=list(verlaat_opties.keys()).index(trap_data.get('verlaat_tijd', 'STANDAARD')),
            key=f"verlaat_{geselecteerde_trap}"
        )

    with col2:
        trap_data['doorgang_uitgang'] = st.number_input(
            "Breedte uitgang trappenhuis (m)",
            min_value=0.5,
            max_value=5.0,
            value=float(trap_data.get('doorgang_uitgang', 0.9)),
            step=0.05,
            key=f"uitgang_{geselecteerde_trap}"
        )

        uitgang_opties = {
            'ENKELE_DEUR_LT135': 'Enkele deur < 135° (110 p/m/min)',
            'DUBBELE_DEUR_LT135': 'Dubbele deur < 135° (90 p/m/min)',
            'ANDERE_DOORGANG': 'Andere doorgang ≥ 135° (135 p/m/min)',
        }
        trap_data['type_uitgang'] = st.selectbox(
            "Type uitgang",
            list(uitgang_opties.keys()),
            format_func=lambda x: uitgang_opties[x],
            index=list(uitgang_opties.keys()).index(trap_data.get('type_uitgang', 'ENKELE_DEUR_LT135')),
            key=f"type_uitgang_{geselecteerde_trap}"
        )

    st.markdown("---")
    st.markdown("### 📐 Verdieping configuratie")

    # Kopieer naar alle functie
    col1, col2 = st.columns([2, 1])
    with col1:
        st.markdown("*Tip: Configureer één verdieping en kopieer naar alle andere.*")
    with col2:
        bron_verd = st.selectbox(
            "Kopieer van verdieping",
            list(trap_data['verdiepingen'].keys()),
            key=f"bron_verd_{geselecteerde_trap}"
        )
        if st.button("📋 Kopieer naar alle", key=f"kopieer_{geselecteerde_trap}"):
            bron_data = trap_data['verdiepingen'][bron_verd]
            for v in trap_data['verdiepingen']:
                if v != bron_verd:
                    trap_data['verdiepingen'][v] = bron_data.copy()
            st.success("Gekopieerd!")
            st.rerun()

    # Verdieping selectie
    geselecteerde_verd = st.selectbox(
        "Selecteer verdieping",
        list(trap_data['verdiepingen'].keys()),
        format_func=lambda x: f"Verdieping {x}" if x != "0" else "Begane grond (0)",
        key=f"verd_select_{geselecteerde_trap}"
    )

    verd_data = trap_data['verdiepingen'][geselecteerde_verd]

    col1, col2, col3 = st.columns(3)

    with col1:
        st.markdown("**Toegang tot trap**")
        verd_data['doorgang_trap'] = st.number_input(
            "Breedte toegangsdeur (m)",
            min_value=0.5,
            max_value=3.0,
            value=float(verd_data.get('doorgang_trap', STANDAARD_DOORGANG_BREEDTE)),
            step=0.05,
            key=f"doorgang_{geselecteerde_trap}_{geselecteerde_verd}"
        )

        doorgang_opties = {
            'ENKELE_DEUR_LT135': 'Enkele deur < 135°',
            'DUBBELE_DEUR_LT135': 'Dubbele deur < 135°',
            'ANDERE_DOORGANG': 'Andere doorgang ≥ 135°',
        }
        verd_data['type_doorgang'] = st.selectbox(
            "Type toegangsdeur",
            list(doorgang_opties.keys()),
            format_func=lambda x: doorgang_opties[x],
            index=list(doorgang_opties.keys()).index(verd_data.get('type_doorgang', 'ENKELE_DEUR_LT135')),
            key=f"type_{geselecteerde_trap}_{geselecteerde_verd}"
        )

    with col2:
        st.markdown("**Trappenhuis**")
        verd_data['trapbreedte'] = st.number_input(
            "Breedte trap (m)",
            min_value=0.8,
            max_value=3.0,
            value=float(verd_data.get('trapbreedte', STANDAARD_TRAPBREEDTE)),
            step=0.05,
            key=f"trapbr_{geselecteerde_trap}_{geselecteerde_verd}"
        )

        verd_data['treden'] = st.number_input(
            "Aantal treden",
            min_value=10,
            max_value=40,
            value=int(verd_data.get('treden', STANDAARD_AANTAL_TREDEN)),
            key=f"treden_{geselecteerde_trap}_{geselecteerde_verd}"
        )

        verd_data['hoogteverschil'] = st.number_input(
            "Hoogteverschil (m)",
            min_value=2.1,
            max_value=4.0,
            value=float(verd_data.get('hoogteverschil', STANDAARD_HOOGTEVERSCHIL)),
            step=0.1,
            key=f"hoogte_{geselecteerde_trap}_{geselecteerde_verd}"
        )

    with col3:
        st.markdown("**Bordessen**")
        verd_data['bordes'] = st.number_input(
            "Oppervlakte bordes (m²)",
            min_value=0.0,
            max_value=20.0,
            value=float(verd_data.get('bordes', STANDAARD_BORDES_OPP)),
            step=0.5,
            key=f"bordes_{geselecteerde_trap}_{geselecteerde_verd}"
        )

        verd_data['tussenbordes'] = st.number_input(
            "Oppervlakte tussenbordessen (m²)",
            min_value=0.0,
            max_value=20.0,
            value=float(verd_data.get('tussenbordes', STANDAARD_TUSSENBORDES_OPP)),
            step=0.5,
            key=f"tussenbordes_{geselecteerde_trap}_{geselecteerde_verd}"
        )

    # Update trap data
    st.session_state.trappen[geselecteerde_trap] = trap_data

    # Navigatie
    st.markdown("---")
    col1, col2, col3 = st.columns([1, 2, 1])
    with col1:
        if st.button("⬅️ Terug", width="stretch"):
            st.session_state.stap = 2
            st.rerun()
    with col3:
        if st.button("🧮 Bereken", type="primary", width="stretch"):
            with st.spinner("Berekening uitvoeren..."):
                voer_berekening_uit()
            st.session_state.stap = 4
            st.rerun()


def voer_berekening_uit():
    """Voer de simulatie en toetsing uit"""
    # Bouw Project object
    project = Project(
        projectnaam=st.session_state.project['projectnaam'],
        aantal_trappen=st.session_state.project['aantal_trappen'],
        aantal_bouwlagen=st.session_state.project['aantal_bouwlagen'],
        laagste_verdieping=st.session_state.project['laagste_verdieping'],
        voorportalen_aanwezig=st.session_state.project['voorportalen'],
    )

    # Update personen
    for v in project.verdiepingen:
        v.aantal_personen = st.session_state.personen.get(str(v.nummer), 0)

    # Bouw Trap objecten
    project.trappen = []

    for trap_naam, trap_data in st.session_state.trappen.items():
        # Bepaal locatie type
        locatie_map = {
            'VEILIGHEIDSVLUCHTROUTE': LocatieType.VEILIGHEIDSVLUCHTROUTE,
            'EXTRA_BESCHERMD': LocatieType.EXTRA_BESCHERMD,
            'BESCHERMD': LocatieType.BESCHERMD,
        }
        locatie = locatie_map.get(trap_data['locatie'], LocatieType.BESCHERMD)

        verlaat_map = {
            'STANDAARD': VerlaatTijdType.STANDAARD,
            'VERLENGD': VerlaatTijdType.VERLENGD,
        }
        verlaat_tijd = verlaat_map.get(trap_data['verlaat_tijd'], VerlaatTijdType.STANDAARD)

        uitgang_map = {
            'ENKELE_DEUR_LT135': DoorgangType.ENKELE_DEUR_LT135,
            'DUBBELE_DEUR_LT135': DoorgangType.DUBBELE_DEUR_LT135,
            'ANDERE_DOORGANG': DoorgangType.ANDERE_DOORGANG,
        }
        type_uitgang = uitgang_map.get(trap_data['type_uitgang'], DoorgangType.ENKELE_DEUR_LT135)

        trap = Trap(
            aanduiding=trap_naam,
            locatie=locatie,
            verlaat_tijd_type=verlaat_tijd,
            doorgang_uitgang_m=trap_data['doorgang_uitgang'],
            type_uitgang=type_uitgang,
        )

        # Bouw verdiepingen
        trap.verdiepingen = []
        for verd_str, verd_data in trap_data['verdiepingen'].items():
            verd_nr = int(verd_str)

            type_doorgang = uitgang_map.get(verd_data['type_doorgang'], DoorgangType.ENKELE_DEUR_LT135)

            tv = TrapVerdieping(
                verdieping=verd_nr,
                doorgang_naar_trap_m=verd_data['doorgang_trap'],
                type_doorgang=type_doorgang,
                oppervlakte_bordes_m2=verd_data['bordes'],
                oppervlakte_tussenbordessen_m2=verd_data['tussenbordes'],
                hoogteverschil_m=verd_data['hoogteverschil'],
                breedte_trap_m=verd_data['trapbreedte'],
                aantal_treden=verd_data['treden'],
            )
            trap.verdiepingen.append(tv)

        trap.sorteer_verdiepingen()
        project.trappen.append(trap)

    # Verdeel personen over trappen
    project.verdeel_personen_over_trappen()

    # Voer simulatie uit
    resultaten = simuleer_alle_trappen(project.trappen)
    toetsingen = toets_alle_resultaten(resultaten, trappen=project.trappen)

    st.session_state.resultaten = resultaten
    st.session_state.toetsingen = toetsingen
    st.session_state.project_obj = project


def render_stap4_resultaten():
    """Render stap 4: Resultaten"""
    st.markdown("## 📊 Stap 4: Resultaten")

    if st.session_state.resultaten is None:
        st.warning("Nog geen berekening uitgevoerd. Ga terug naar stap 3 en klik op 'Bereken'.")
        if st.button("⬅️ Terug naar configuratie"):
            st.session_state.stap = 3
            st.rerun()
        return

    resultaten = st.session_state.resultaten
    toetsingen = st.session_state.toetsingen

    # Overzicht metrics
    st.markdown("### 📈 Samenvatting")

    cols = st.columns(len(resultaten))
    for i, (trap_naam, res) in enumerate(resultaten.items()):
        toets = toetsingen[trap_naam]
        with cols[i]:
            status = "✅" if toets.alle_criteria_voldaan else "❌"
            st.metric(
                label=f"{trap_naam} {status}",
                value=f"{res.totaal_personen} pers",
                delta=f"Klaar in {res.voltooide_ontruiming_min or res.ontruimingstijd_min:.1f} min"
            )

    # Tabs voor verschillende views
    tab1, tab2, tab3, tab4, tab5 = st.tabs(["📊 Grafieken", "✅ Normtoetsing", "🤖 AI Advies", "📋 Details", "📥 Rapport"])

    with tab1:
        render_grafieken(resultaten)

    with tab2:
        render_normtoetsing(toetsingen)

    with tab3:
        render_ai_advies(resultaten, toetsingen)

    with tab4:
        render_details(resultaten)

    with tab5:
        render_rapport_export(resultaten, toetsingen)

    # Navigatie
    st.markdown("---")
    col1, col2, col3 = st.columns([1, 2, 1])
    with col1:
        if st.button("⬅️ Terug", width="stretch"):
            st.session_state.stap = 3
            st.rerun()
    with col2:
        if st.button("🔄 Herbereken", width="stretch"):
            with st.spinner("Berekening uitvoeren..."):
                voer_berekening_uit()
            st.rerun()


def render_grafieken(resultaten):
    """Render grafieken van de resultaten"""
    st.markdown("### 📈 Ontruimingsverloop")

    # Maak Plotly figuur
    fig = go.Figure()

    colors = px.colors.qualitative.Set2

    for i, (trap_naam, res) in enumerate(resultaten.items()):
        tijden = [ts.tijd_min for ts in res.tijdstappen]
        buiten = [ts.cumulatief_buiten for ts in res.tijdstappen]

        fig.add_trace(go.Scatter(
            x=tijden,
            y=buiten,
            mode='lines+markers',
            name=trap_naam,
            line=dict(color=colors[i % len(colors)], width=2),
            marker=dict(size=4)
        ))

        # Voeg horizontale lijn toe voor totaal personen
        fig.add_hline(
            y=res.totaal_personen,
            line_dash="dash",
            line_color=colors[i % len(colors)],
            annotation_text=f"{trap_naam}: {res.totaal_personen}",
            annotation_position="right"
        )

    fig.update_layout(
        title="Cumulatief aantal personen buiten",
        xaxis_title="Tijd (minuten)",
        yaxis_title="Aantal personen",
        legend=dict(orientation="h", yanchor="bottom", y=1.02, xanchor="right", x=1),
        hovermode="x unified"
    )

    st.plotly_chart(fig, width="stretch")

    # Tweede grafiek: personen in trap
    st.markdown("### 🚶 Personen in trappenhuis")

    fig2 = go.Figure()

    for i, (trap_naam, res) in enumerate(resultaten.items()):
        tijden = [ts.tijd_min for ts in res.tijdstappen]
        in_trap = [ts.totaal_in_trap for ts in res.tijdstappen]

        fig2.add_trace(go.Scatter(
            x=tijden,
            y=in_trap,
            mode='lines',
            name=trap_naam,
            fill='tozeroy',
            line=dict(color=colors[i % len(colors)])
        ))

    fig2.update_layout(
        title="Personen in trappenhuis over tijd",
        xaxis_title="Tijd (minuten)",
        yaxis_title="Aantal personen in trap",
        legend=dict(orientation="h", yanchor="bottom", y=1.02, xanchor="right", x=1)
    )

    st.plotly_chart(fig2, width="stretch")


def render_normtoetsing(toetsingen):
    """Render normtoetsing resultaten"""
    st.markdown("### ✅ BBL Art. 4.81 Normtoetsing")

    for trap_naam, toets in toetsingen.items():
        with st.expander(f"**{trap_naam}** - {'✅ Voldoet' if toets.alle_criteria_voldaan else '❌ Voldoet niet'}", expanded=True):

            for criterium in toets.criteria:
                if criterium.status == ToetsStatus.VOLDOET:
                    st.markdown(f"""
                    <div class="success-box">
                        <strong>✅ {criterium.naam}</strong><br>
                        {criterium.beschrijving}<br>
                        <em>Eis: {criterium.eis} {criterium.eenheid} | Berekend: {criterium.berekend:.1f} {criterium.eenheid}</em><br>
                        {criterium.toelichting}
                    </div>
                    """, unsafe_allow_html=True)
                else:
                    st.markdown(f"""
                    <div class="error-box">
                        <strong>❌ {criterium.naam}</strong><br>
                        {criterium.beschrijving}<br>
                        <em>Eis: {criterium.eis} {criterium.eenheid} | Berekend: {criterium.berekend:.1f} {criterium.eenheid}</em><br>
                        {criterium.toelichting}
                    </div>
                    """, unsafe_allow_html=True)


def get_ai_advies(resultaten, toetsingen):
    """Genereer AI advies met Google Gemini"""
    try:
        # Haal API key uit secrets
        api_key = st.secrets.get("GOOGLE_API_KEY")
        if not api_key:
            return None, "API key niet geconfigureerd. Voeg GOOGLE_API_KEY toe aan Streamlit Secrets."

        genai.configure(api_key=api_key)
        # Probeer verschillende modellen als fallback
        model = genai.GenerativeModel('gemini-2.5-flash')

        # Bouw context voor AI
        context = "Je bent een expert in brandveiligheid en evacuatieberekeningen volgens het Nederlandse Besluit bouwwerken leefomgeving (BBL).\n\n"
        context += "Analyseer de volgende evacuatieberekening en geef praktisch advies:\n\n"

        for trap_naam, res in resultaten.items():
            toets = toetsingen[trap_naam]
            context += f"## {trap_naam}\n"
            context += f"- Totaal personen: {res.totaal_personen}\n"
            context += f"- Ontruimingstijd beschikbaar: {res.ontruimingstijd_min} minuten\n"
            context += f"- Werkelijke ontruimingstijd: {res.voltooide_ontruiming_min or 'niet voltooid'} minuten\n"
            context += f"- Status: {'VOLDOET' if toets.alle_criteria_voldaan else 'VOLDOET NIET'}\n\n"

            for criterium in toets.criteria:
                status = "✅" if criterium.voldoet else "❌"
                context += f"  {status} {criterium.naam}: eis={criterium.eis}, berekend={criterium.berekend:.1f}\n"
                if not criterium.voldoet:
                    context += f"     Probleem: {criterium.toelichting}\n"
            context += "\n"

        # Voeg trap configuratie toe
        if 'trappen' in st.session_state:
            context += "## Trap Configuratie\n"
            for trap_naam, trap_data in st.session_state.trappen.items():
                context += f"### {trap_naam}\n"
                context += f"- Uitgang breedte: {trap_data.get('doorgang_uitgang', 'onbekend')}m\n"
                context += f"- Type vluchtroute: {trap_data.get('locatie', 'onbekend')}\n"

        prompt = context + """

Geef concreet advies in het Nederlands:
1. Samenvatting: voldoet het ontwerp aan BBL Art. 4.81?
2. Als er problemen zijn: wat zijn de specifieke knelpunten?
3. Concrete oplossingen met geschatte impact (bijv. "verbreed deur van 0.85m naar 1.2m = +40% capaciteit")
4. Prioritering: wat heeft de meeste impact?

Houd het praktisch en beknopt (max 300 woorden)."""

        response = model.generate_content(prompt)
        return response.text, None

    except Exception as e:
        return None, f"Fout bij AI analyse: {str(e)}"


def render_ai_advies(resultaten, toetsingen):
    """Render AI advies sectie"""
    st.markdown("### 🤖 AI Advies")

    st.markdown("""
    <div class="info-box">
        Krijg automatisch advies van AI over je evacuatieberekening.
        De AI analyseert de resultaten en geeft concrete verbetervoorstellen.
    </div>
    """, unsafe_allow_html=True)

    # Check of er al advies is
    if 'ai_advies' not in st.session_state:
        st.session_state.ai_advies = None

    col1, col2 = st.columns([1, 3])
    with col1:
        if st.button("🤖 Genereer AI Advies", type="primary"):
            with st.spinner("AI analyseert de resultaten..."):
                advies, error = get_ai_advies(resultaten, toetsingen)
                if error:
                    st.error(error)
                else:
                    st.session_state.ai_advies = advies
                    st.rerun()

    # Toon advies als beschikbaar
    if st.session_state.ai_advies:
        st.markdown("---")
        st.markdown("#### 💡 AI Analyse")
        st.markdown(st.session_state.ai_advies)

        st.markdown("---")
        st.caption("*Dit advies is gegenereerd door AI en dient ter indicatie. Raadpleeg altijd een gekwalificeerd adviseur voor definitieve beslissingen.*")


def maak_evacuatie_grafiek(resultaten):
    """Maak de cumulatieve evacuatie grafiek"""
    fig = go.Figure()
    colors = px.colors.qualitative.Set2

    for i, (trap_naam, res) in enumerate(resultaten.items()):
        tijden = [ts.tijd_min for ts in res.tijdstappen]
        buiten = [ts.cumulatief_buiten for ts in res.tijdstappen]

        fig.add_trace(go.Scatter(
            x=tijden,
            y=buiten,
            mode='lines+markers',
            name=trap_naam,
            line=dict(color=colors[i % len(colors)], width=2),
            marker=dict(size=4)
        ))

        fig.add_hline(
            y=res.totaal_personen,
            line_dash="dash",
            line_color=colors[i % len(colors)],
            annotation_text=f"{trap_naam}: {res.totaal_personen}",
            annotation_position="right"
        )

    fig.update_layout(
        title="Cumulatief aantal personen buiten",
        xaxis_title="Tijd (minuten)",
        yaxis_title="Aantal personen",
        legend=dict(orientation="h", yanchor="bottom", y=1.02, xanchor="right", x=1),
        width=700,
        height=400
    )
    return fig


def maak_trap_bezetting_grafiek(resultaten):
    """Maak de trap bezetting grafiek"""
    fig = go.Figure()
    colors = px.colors.qualitative.Set2

    for i, (trap_naam, res) in enumerate(resultaten.items()):
        tijden = [ts.tijd_min for ts in res.tijdstappen]
        in_trap = [ts.totaal_in_trap for ts in res.tijdstappen]

        fig.add_trace(go.Scatter(
            x=tijden,
            y=in_trap,
            mode='lines',
            name=trap_naam,
            fill='tozeroy',
            line=dict(color=colors[i % len(colors)])
        ))

    fig.update_layout(
        title="Personen in trappenhuis over tijd",
        xaxis_title="Tijd (minuten)",
        yaxis_title="Aantal personen in trap",
        legend=dict(orientation="h", yanchor="bottom", y=1.02, xanchor="right", x=1),
        width=700,
        height=400
    )
    return fig


def genereer_pdf_rapport(resultaten, toetsingen):
    """Genereer een PDF rapport van de evacuatieberekening"""
    pdf = FPDF()
    pdf.add_page()
    pdf.set_auto_page_break(auto=True, margin=15)

    # Titel
    pdf.set_font('Helvetica', 'B', 20)
    pdf.cell(0, 15, 'Evacuatie Calculator Rapport', ln=True, align='C')
    pdf.set_font('Helvetica', '', 10)
    pdf.cell(0, 8, 'BBL Art. 4.80/4.81 - Opvang- en doorstroomcapaciteit', ln=True, align='C')
    pdf.ln(10)

    # Project informatie
    pdf.set_font('Helvetica', 'B', 14)
    pdf.cell(0, 10, '1. Projectgegevens', ln=True)
    pdf.set_font('Helvetica', '', 11)

    project = st.session_state.project
    pdf.cell(0, 7, f"Projectnaam: {project.get('projectnaam', '-')}", ln=True)
    pdf.cell(0, 7, f"Datum: {project.get('datum', '-')}", ln=True)
    pdf.cell(0, 7, f"Aantal bouwlagen: {project.get('aantal_bouwlagen', '-')}", ln=True)
    pdf.cell(0, 7, f"Laagste verdieping: {project.get('laagste_verdieping', 0)}", ln=True)
    pdf.cell(0, 7, f"Aantal trappenhuizen: {project.get('aantal_trappen', '-')}", ln=True)
    pdf.cell(0, 7, f"Voorportalen: {'Ja' if project.get('voorportalen') else 'Nee'}", ln=True)
    pdf.ln(5)

    # Personen per verdieping
    pdf.set_font('Helvetica', 'B', 14)
    pdf.cell(0, 10, '2. Personen per verdieping', ln=True)
    pdf.set_font('Helvetica', '', 11)

    totaal_personen = 0
    for verd, aantal in sorted(st.session_state.personen.items(), key=lambda x: int(x[0]), reverse=True):
        pdf.cell(0, 7, f"Verdieping {verd}: {aantal} personen", ln=True)
        totaal_personen += aantal
    pdf.set_font('Helvetica', 'B', 11)
    pdf.cell(0, 7, f"Totaal: {totaal_personen} personen", ln=True)
    pdf.ln(5)

    # Trap configuratie
    pdf.set_font('Helvetica', 'B', 14)
    pdf.cell(0, 10, '3. Trapconfiguratie', ln=True)

    for trap_naam, trap_data in st.session_state.trappen.items():
        pdf.set_font('Helvetica', 'B', 12)
        pdf.cell(0, 8, trap_naam, ln=True)
        pdf.set_font('Helvetica', '', 11)
        pdf.cell(0, 6, f"  Trapbreedte: {trap_data.get('trapbreedte', '-')} m", ln=True)
        pdf.cell(0, 6, f"  Uitgang breedte: {trap_data.get('doorgang_uitgang', '-')} m", ln=True)
        pdf.cell(0, 6, f"  Locatie type: {trap_data.get('locatie', '-')}", ln=True)
    pdf.ln(5)

    # Resultaten per trap
    pdf.add_page()
    pdf.set_font('Helvetica', 'B', 14)
    pdf.cell(0, 10, '4. Resultaten', ln=True)

    for trap_naam, res in resultaten.items():
        toets = toetsingen[trap_naam]
        status = "VOLDOET" if toets.alle_criteria_voldaan else "VOLDOET NIET"

        pdf.set_font('Helvetica', 'B', 12)
        pdf.cell(0, 8, f"{trap_naam} - {status}", ln=True)
        pdf.set_font('Helvetica', '', 11)
        pdf.cell(0, 6, f"  Totaal personen: {res.totaal_personen}", ln=True)
        pdf.cell(0, 6, f"  Beschikbare ontruimingstijd: {res.ontruimingstijd_min} min", ln=True)
        pdf.cell(0, 6, f"  Werkelijke ontruimingstijd: {res.voltooide_ontruiming_min or 'niet voltooid'} min", ln=True)
        pdf.ln(3)

        # Toetsingscriteria
        pdf.set_font('Helvetica', 'B', 11)
        pdf.cell(0, 7, "  Toetsingscriteria BBL Art. 4.81:", ln=True)
        pdf.set_font('Helvetica', '', 10)

        for criterium in toets.criteria:
            status_icon = "[OK]" if criterium.voldoet else "[X]"
            pdf.cell(0, 6, f"    {status_icon} {criterium.naam}: eis={criterium.eis}, berekend={criterium.berekend:.1f} {criterium.eenheid}", ln=True)
            if not criterium.voldoet:
                pdf.set_text_color(180, 0, 0)
                pdf.cell(0, 5, f"        {criterium.toelichting}", ln=True)
                pdf.set_text_color(0, 0, 0)
        pdf.ln(5)

    # Grafieken
    pdf.add_page()
    pdf.set_font('Helvetica', 'B', 14)
    pdf.cell(0, 10, '5. Grafieken', ln=True)
    pdf.ln(5)

    try:
        # Evacuatie grafiek
        fig1 = maak_evacuatie_grafiek(resultaten)
        img1_bytes = fig1.to_image(format="png", scale=2)
        pdf.image(io.BytesIO(img1_bytes), x=10, w=190)
        pdf.ln(10)

        # Trap bezetting grafiek
        fig2 = maak_trap_bezetting_grafiek(resultaten)
        img2_bytes = fig2.to_image(format="png", scale=2)
        pdf.image(io.BytesIO(img2_bytes), x=10, w=190)
    except Exception as e:
        pdf.set_font('Helvetica', 'I', 10)
        pdf.cell(0, 10, f"Grafieken konden niet worden gegenereerd: {str(e)}", ln=True)

    # AI Advies indien beschikbaar
    if st.session_state.get('ai_advies'):
        pdf.add_page()
        pdf.set_font('Helvetica', 'B', 14)
        pdf.cell(0, 10, '6. AI Advies', ln=True)
        pdf.set_font('Helvetica', '', 10)

        # Split advies in regels en voeg toe
        advies_text = st.session_state.ai_advies
        pdf.multi_cell(0, 6, advies_text)
        pdf.ln(5)
        pdf.set_font('Helvetica', 'I', 9)
        pdf.cell(0, 6, "Dit advies is gegenereerd door AI en dient ter indicatie.", ln=True)

    # Footer
    pdf.ln(10)
    pdf.set_font('Helvetica', 'I', 9)
    pdf.cell(0, 6, f"Gegenereerd met Evacuatie Calculator - {date.today()}", ln=True)

    return bytes(pdf.output())


def genereer_word_rapport(resultaten, toetsingen):
    """Genereer een Word rapport van de evacuatieberekening"""
    doc = Document()

    # Titel
    title = doc.add_heading('Evacuatie Calculator Rapport', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph('BBL Art. 4.80/4.81 - Opvang- en doorstroomcapaciteit')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Project informatie
    doc.add_heading('1. Projectgegevens', level=1)
    project = st.session_state.project

    table = doc.add_table(rows=6, cols=2)
    table.style = 'Table Grid'

    cells = [
        ('Projectnaam', project.get('projectnaam', '-')),
        ('Datum', str(project.get('datum', '-'))),
        ('Aantal bouwlagen', str(project.get('aantal_bouwlagen', '-'))),
        ('Laagste verdieping', str(project.get('laagste_verdieping', 0))),
        ('Aantal trappenhuizen', str(project.get('aantal_trappen', '-'))),
        ('Voorportalen', 'Ja' if project.get('voorportalen') else 'Nee'),
    ]

    for i, (label, value) in enumerate(cells):
        table.rows[i].cells[0].text = label
        table.rows[i].cells[1].text = value

    doc.add_paragraph()

    # Personen per verdieping
    doc.add_heading('2. Personen per verdieping', level=1)

    personen_sorted = sorted(st.session_state.personen.items(), key=lambda x: int(x[0]), reverse=True)
    table = doc.add_table(rows=len(personen_sorted) + 1, cols=2)
    table.style = 'Table Grid'

    table.rows[0].cells[0].text = 'Verdieping'
    table.rows[0].cells[1].text = 'Aantal personen'

    totaal = 0
    for i, (verd, aantal) in enumerate(personen_sorted, 1):
        table.rows[i].cells[0].text = f"Verdieping {verd}"
        table.rows[i].cells[1].text = str(aantal)
        totaal += aantal

    doc.add_paragraph(f"Totaal: {totaal} personen", style='Intense Quote')

    # Trap configuratie
    doc.add_heading('3. Trapconfiguratie', level=1)

    for trap_naam, trap_data in st.session_state.trappen.items():
        doc.add_heading(trap_naam, level=2)

        table = doc.add_table(rows=3, cols=2)
        table.style = 'Table Grid'

        table.rows[0].cells[0].text = 'Trapbreedte'
        table.rows[0].cells[1].text = f"{trap_data.get('trapbreedte', '-')} m"
        table.rows[1].cells[0].text = 'Uitgang breedte'
        table.rows[1].cells[1].text = f"{trap_data.get('doorgang_uitgang', '-')} m"
        table.rows[2].cells[0].text = 'Locatie type'
        table.rows[2].cells[1].text = str(trap_data.get('locatie', '-'))

    doc.add_page_break()

    # Resultaten
    doc.add_heading('4. Resultaten', level=1)

    for trap_naam, res in resultaten.items():
        toets = toetsingen[trap_naam]
        status = "VOLDOET" if toets.alle_criteria_voldaan else "VOLDOET NIET"

        doc.add_heading(f"{trap_naam} - {status}", level=2)

        # Samenvatting
        table = doc.add_table(rows=3, cols=2)
        table.style = 'Table Grid'

        table.rows[0].cells[0].text = 'Totaal personen'
        table.rows[0].cells[1].text = str(res.totaal_personen)
        table.rows[1].cells[0].text = 'Beschikbare ontruimingstijd'
        table.rows[1].cells[1].text = f"{res.ontruimingstijd_min} min"
        table.rows[2].cells[0].text = 'Werkelijke ontruimingstijd'
        table.rows[2].cells[1].text = f"{res.voltooide_ontruiming_min or 'niet voltooid'} min"

        doc.add_paragraph()
        doc.add_paragraph('Toetsingscriteria BBL Art. 4.81:', style='Heading 3')

        # Criteria tabel
        table = doc.add_table(rows=len(toets.criteria) + 1, cols=4)
        table.style = 'Table Grid'

        headers = ['Status', 'Criterium', 'Eis', 'Berekend']
        for i, header in enumerate(headers):
            table.rows[0].cells[i].text = header

        for i, criterium in enumerate(toets.criteria, 1):
            status_icon = "OK" if criterium.voldoet else "NIET OK"
            table.rows[i].cells[0].text = status_icon
            table.rows[i].cells[1].text = criterium.naam
            table.rows[i].cells[2].text = f"{criterium.eis} {criterium.eenheid}"
            table.rows[i].cells[3].text = f"{criterium.berekend:.1f} {criterium.eenheid}"

        doc.add_paragraph()

    # Grafieken
    doc.add_page_break()
    doc.add_heading('5. Grafieken', level=1)

    try:
        # Evacuatie grafiek
        fig1 = maak_evacuatie_grafiek(resultaten)
        img1_bytes = fig1.to_image(format="png", scale=2)
        doc.add_paragraph("Ontruimingsverloop - Cumulatief aantal personen buiten:")
        doc.add_picture(io.BytesIO(img1_bytes), width=Inches(6))
        doc.add_paragraph()

        # Trap bezetting grafiek
        fig2 = maak_trap_bezetting_grafiek(resultaten)
        img2_bytes = fig2.to_image(format="png", scale=2)
        doc.add_paragraph("Bezetting trappenhuis over tijd:")
        doc.add_picture(io.BytesIO(img2_bytes), width=Inches(6))
    except Exception as e:
        doc.add_paragraph(f"Grafieken konden niet worden gegenereerd: {str(e)}")

    # Tijdstap data per trap
    doc.add_page_break()
    doc.add_heading('6. Gedetailleerde tijdstapdata', level=1)

    for trap_naam, res in resultaten.items():
        doc.add_heading(trap_naam, level=2)

        # Beperk tot eerste 20 tijdstappen voor leesbaarheid
        tijdstappen = res.tijdstappen[:20]
        table = doc.add_table(rows=len(tijdstappen) + 1, cols=5)
        table.style = 'Table Grid'

        headers = ['Tijd (min)', 'Naar buiten', 'Cum. buiten', 'In trap', 'Op verd.']
        for i, header in enumerate(headers):
            table.rows[0].cells[i].text = header

        for i, ts in enumerate(tijdstappen, 1):
            table.rows[i].cells[0].text = f"{ts.tijd_min:.1f}"
            table.rows[i].cells[1].text = f"{ts.naar_buiten:.0f}"
            table.rows[i].cells[2].text = f"{ts.cumulatief_buiten:.0f}"
            table.rows[i].cells[3].text = f"{ts.totaal_in_trap:.0f}"
            table.rows[i].cells[4].text = f"{ts.totaal_op_verdiepingen:.0f}"

        if len(res.tijdstappen) > 20:
            doc.add_paragraph(f"... en {len(res.tijdstappen) - 20} meer tijdstappen")

        doc.add_paragraph()

    # AI Advies indien beschikbaar
    if st.session_state.get('ai_advies'):
        doc.add_page_break()
        doc.add_heading('7. AI Advies', level=1)
        doc.add_paragraph(st.session_state.ai_advies)
        doc.add_paragraph()
        disclaimer = doc.add_paragraph("Dit advies is gegenereerd door AI en dient ter indicatie. Raadpleeg altijd een gekwalificeerd adviseur voor definitieve beslissingen.")
        disclaimer.italic = True

    # Footer
    doc.add_paragraph()
    footer = doc.add_paragraph(f"Gegenereerd met Evacuatie Calculator - {date.today()}")
    footer.italic = True

    # Save to bytes
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


def render_rapport_export(resultaten, toetsingen):
    """Render de rapport export sectie"""
    st.markdown("### 📥 Rapport Exporteren")

    st.markdown("""
    <div class="info-box">
        Download een volledig rapport van de evacuatieberekening als PDF of Word document.
        Het rapport bevat alle projectgegevens, configuratie, resultaten en toetsingen.
    </div>
    """, unsafe_allow_html=True)

    col1, col2 = st.columns(2)

    with col1:
        st.markdown("#### 📄 PDF Rapport")
        st.markdown("Geschikt voor delen en archiveren.")

        if st.button("📄 Genereer PDF", type="primary", key="gen_pdf"):
            with st.spinner("PDF genereren..."):
                try:
                    pdf_bytes = genereer_pdf_rapport(resultaten, toetsingen)
                    st.session_state.pdf_rapport = pdf_bytes
                    st.success("PDF gegenereerd!")
                except Exception as e:
                    st.error(f"Fout bij genereren PDF: {str(e)}")

        if st.session_state.get('pdf_rapport'):
            st.download_button(
                label="⬇️ Download PDF",
                data=st.session_state.pdf_rapport,
                file_name=f"{st.session_state.project['projectnaam']}_rapport.pdf",
                mime="application/pdf"
            )

    with col2:
        st.markdown("#### 📝 Word Rapport")
        st.markdown("Geschikt voor bewerken en aanpassen.")

        if st.button("📝 Genereer Word", type="primary", key="gen_word"):
            with st.spinner("Word document genereren..."):
                try:
                    word_bytes = genereer_word_rapport(resultaten, toetsingen)
                    st.session_state.word_rapport = word_bytes
                    st.success("Word document gegenereerd!")
                except Exception as e:
                    st.error(f"Fout bij genereren Word: {str(e)}")

        if st.session_state.get('word_rapport'):
            st.download_button(
                label="⬇️ Download Word",
                data=st.session_state.word_rapport,
                file_name=f"{st.session_state.project['projectnaam']}_rapport.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

    st.markdown("---")
    st.markdown("#### 📋 Wat bevat het rapport?")
    st.markdown("""
    - **Projectgegevens**: naam, nummer, locatie, datum
    - **Gebouwconfiguratie**: verdiepingen, personen per verdieping
    - **Trapconfiguratie**: afmetingen, type vluchtroute
    - **Resultaten**: ontruimingstijden per trap
    - **BBL Toetsing**: alle criteria met status
    - **Tijdstapdata**: gedetailleerd verloop evacuatie
    - **AI Advies**: indien gegenereerd
    """)


def render_details(resultaten):
    """Render gedetailleerde resultaten"""
    st.markdown("### 📋 Gedetailleerde tijdstapdata")

    for trap_naam, res in resultaten.items():
        with st.expander(f"**{trap_naam}** - Tijdstap data"):
            # Maak tabel data
            data = []
            for ts in res.tijdstappen:
                row = {
                    'Tijd (min)': ts.tijd_min,
                    'Naar buiten': f"{ts.naar_buiten:.0f}",
                    'Cumulatief buiten': f"{ts.cumulatief_buiten:.0f}",
                    'In trap': f"{ts.totaal_in_trap:.0f}",
                    'Op verdiepingen': f"{ts.totaal_op_verdiepingen:.0f}",
                }
                data.append(row)

            st.dataframe(data, width="stretch", height=400)

            # Download optie
            df = pd.DataFrame(data)
            csv = df.to_csv(index=False)
            st.download_button(
                label="⬇️ Download CSV",
                data=csv,
                file_name=f"{trap_naam}_data.csv",
                mime="text/csv"
            )


def main():
    """Hoofdfunctie"""
    init_session_state()
    render_header()
    render_sidebar()

    # Render huidige stap
    if st.session_state.stap == 1:
        render_stap1_project()
    elif st.session_state.stap == 2:
        render_stap2_personen()
    elif st.session_state.stap == 3:
        render_stap3_trappen()
    elif st.session_state.stap == 4:
        render_stap4_resultaten()


if __name__ == "__main__":
    main()
